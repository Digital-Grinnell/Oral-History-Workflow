#!/usr/bin/env python3
"""
Whisper Transcription Module
Transcribes audio files using OpenAI Whisper and creates DOCX in MS Word format
"""

import whisper
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import warnings

# Suppress FP16 warning if no GPU
warnings.filterwarnings("ignore", message="FP16 is not supported on CPU")


def format_timestamp(seconds):
    """Convert seconds to HH:MM:SS format."""
    hours = int(seconds // 3600)
    minutes = int((seconds % 3600) // 60)
    secs = int(seconds % 60)
    
    if hours > 0:
        return f"{hours}:{minutes:02d}:{secs:02d}"
    else:
        return f"{minutes}:{secs:02d}"


def transcribe_audio(audio_path, model_size="base", language=None, progress_callback=None):
    """
    Transcribe audio file using Whisper.
    
    Args:
        audio_path: Path to audio file (mp3, wav, etc.)
        model_size: Whisper model size (tiny, base, small, medium, large)
        language: Language code (e.g., 'en') or None for auto-detect
        progress_callback: Optional function to call with progress updates
        
    Returns:
        dict with 'text', 'segments', and 'language'
    """
    if progress_callback:
        progress_callback(f"Loading Whisper model ({model_size})...")
    
    # Load model
    model = whisper.load_model(model_size)
    
    if progress_callback:
        progress_callback(f"Transcribing audio file...")
    
    # Transcribe
    result = model.transcribe(
        str(audio_path),
        language=language,
        verbose=False,
        word_timestamps=False
    )
    
    if progress_callback:
        progress_callback("Transcription complete!")
    
    return result


def create_docx_from_transcription(result, output_path, audio_filename):
    """
    Create a DOCX file from Whisper transcription in MS Word format.
    
    Args:
        result: Whisper transcription result dict
        output_path: Path to save DOCX file
        audio_filename: Name of original audio file
    """
    doc = Document()
    
    # Add title
    title = doc.add_heading(f'Transcript: {audio_filename}', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add metadata
    doc.add_paragraph(f"Language: {result.get('language', 'Unknown')}")
    doc.add_paragraph("")  # Blank line
    
    # Add transcription with timestamps
    # Format similar to MS Word Online transcription
    for segment in result.get('segments', []):
        timestamp = format_timestamp(segment['start'])
        text = segment['text'].strip()
        
        # Create paragraph with timestamp
        p = doc.add_paragraph()
        
        # Add speaker (generic for now)
        speaker_run = p.add_run(f"Speaker {timestamp}")
        speaker_run.bold = True
        speaker_run.font.color.rgb = RGBColor(0, 0, 0)
        
        # Add line break
        p.add_run("\n")
        
        # Add text
        text_run = p.add_run(text)
        text_run.font.size = Pt(11)
        
        # Add spacing after paragraph
        p.paragraph_format.space_after = Pt(6)
    
    # Save document
    doc.save(str(output_path))
    return True


def transcribe_and_create_docx(audio_path, output_docx=None, model_size="base", 
                                language=None, progress_callback=None):
    """
    Complete workflow: transcribe audio and create DOCX file.
    
    Args:
        audio_path: Path to audio file
        output_docx: Path to output DOCX (defaults to audio_file.docx)
        model_size: Whisper model size (tiny, base, small, medium, large)
        language: Language code or None for auto-detect
        progress_callback: Optional function for progress updates
        
    Returns:
        tuple: (success: bool, message: str, docx_path: Path)
    """
    try:
        audio_path = Path(audio_path)
        
        if not audio_path.exists():
            return False, f"Audio file not found: {audio_path}", None
        
        # Determine output path
        if output_docx is None:
            output_docx = audio_path.with_suffix('.docx')
        else:
            output_docx = Path(output_docx)
        
        # Transcribe
        result = transcribe_audio(
            audio_path, 
            model_size=model_size, 
            language=language,
            progress_callback=progress_callback
        )
        
        if progress_callback:
            progress_callback("Creating DOCX file...")
        
        # Create DOCX
        create_docx_from_transcription(result, output_docx, audio_path.name)
        
        return True, f"Transcription saved to {output_docx}", output_docx
        
    except Exception as e:
        return False, f"Error during transcription: {str(e)}", None


if __name__ == "__main__":
    import sys
    
    if len(sys.argv) < 2:
        print("Usage: python transcribe_whisper.py <audio_file> [model_size]")
        print("Model sizes: tiny, base (default), small, medium, large")
        sys.exit(1)
    
    audio_file = sys.argv[1]
    model_size = sys.argv[2] if len(sys.argv) > 2 else "base"
    
    def print_progress(msg):
        print(f"[Whisper] {msg}")
    
    success, message, output = transcribe_and_create_docx(
        audio_file,
        model_size=model_size,
        progress_callback=print_progress
    )
    
    if success:
        print(f"✅ {message}")
        sys.exit(0)
    else:
        print(f"❌ {message}")
        sys.exit(1)
